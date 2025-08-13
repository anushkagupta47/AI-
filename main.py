"""
AI Resume Optimiser & Generator (Task 2)
-------------------------------------------------
A Streamlit app that:
- Accepts a candidate resume (PDF/DOCX) and a target job description (text or file)
- Parses & analyses content
- Extracts/aligns keywords
- Generates an ATS-friendly optimised resume (DOCX)
- Produces a concise change log (keywords, formatting, ATS checks)

How to run locally:
1) Create & activate a virtual environment (Python 3.10+ recommended)
2) pip install -r requirements.txt  (see inline list below)
3) streamlit run app.py

Inline requirements (copy to requirements.txt):
-------------------------------------------------
streamlit
pdfplumber
python-docx
scikit-learn
nltk
rapidfuzz

Optional (uncomment spaCy parts if you install):
spacy
# python -m spacy download en_core_web_sm

Notes:
- This app focuses on ATS-friendly text (no tables/images, simple structure)
- Supports PDF and DOCX. Convert old .doc files to .docx before upload.
- No external APIs are required.
"""

import io
import re
import time
from dataclasses import dataclass
from typing import Dict, List, Tuple

import streamlit as st
from rapidfuzz import fuzz
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.pipeline import Pipeline
import nltk

# Ensure NLTK stopwords are available (safe to call; downloads once and caches)
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

from nltk.corpus import stopwords
STOPWORDS = set(stopwords.words('english'))

# Lazy imports for heavy libraries
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None

# ----------------------------- Utilities ----------------------------- #

SECTION_HEADERS = [
    'summary', 'professional summary', 'profile', 'objective',
    'skills', 'key skills', 'technical skills', 'core competencies',
    'experience', 'work experience', 'professional experience', 'employment',
    'projects', 'relevant projects', 'academic projects',
    'education', 'certifications', 'achievements', 'publications',
]

HEADER_PATTERN = re.compile(r"^(?:" + r"|".join([re.escape(h) for h in SECTION_HEADERS]) + r")[\s:]*$", re.I)
BULLET_PATTERN = re.compile(r"^[\-•\u2022\*\>>]+\s+")
CONTACT_PATTERN = re.compile(
    r"(?P<email>[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})|"
    r"(?P<phone>\+?\d[\d\s\-()]{7,}\d)|"
    r"(?P<linkedin>linkedin\.com/in/[A-Za-z0-9\-_/]+)|"
    r"(?P<github>github\.com/[A-Za-z0-9\-_/]+)",
    re.I
)

@dataclass
class ParsedResume:
    contact_block: str
    sections: Dict[str, List[str]]  # section_name -> list of lines
    raw_text: str


def read_pdf(file_bytes: bytes) -> str:
    if pdfplumber is None:
        raise RuntimeError("pdfplumber is not installed.")
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def read_docx(file_bytes: bytes) -> str:
    if Document is None:
        raise RuntimeError("python-docx is not installed.")
    doc = Document(io.BytesIO(file_bytes))
    lines = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            lines.append(txt)
    # Also parse simple tables as text (ATS often ignores tables; we flatten them)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                lines.append(" | ".join(row_text))
    return "\n".join(lines)


def clean_lines(text: str) -> List[str]:
    lines = [re.sub(r"\s+", " ", ln).strip() for ln in text.splitlines()]
    return [ln for ln in lines if ln]


def split_sections(lines: List[str]) -> ParsedResume:
    sections: Dict[str, List[str]] = {}
    current = "misc"
    sections[current] = []
    contact_lines: List[str] = []

    # Heuristic: first 6-10 lines often contain contact info
    for i, ln in enumerate(lines):
        if i < 12 and CONTACT_PATTERN.search(ln):
            contact_lines.append(ln)

        if HEADER_PATTERN.match(ln.lower()):
            current = ln.strip().lower()
            current = re.sub(r"\s+", " ", current)
            sections[current] = []
            continue

        # Normalize bullets to '-' for consistency
        norm = BULLET_PATTERN.sub("- ", ln)
        sections.setdefault(current, []).append(norm)

    contact_block = " | ".join(sorted(set(contact_lines)))
    return ParsedResume(contact_block=contact_block, sections=sections, raw_text="\n".join(lines))


def tfidf_keywords(text: str, top_k: int = 40) -> List[str]:
    # Simple TF-IDF keyword extraction (unigrams & bigrams)
    corpus = [text]
    vectorizer = TfidfVectorizer(
        stop_words='english',
        ngram_range=(1, 2),
        max_features=2000,
        min_df=1
    )
    X = vectorizer.fit_transform(corpus)
    scores = X.toarray()[0]
    terms = vectorizer.get_feature_names_out()
    scored = sorted(zip(terms, scores), key=lambda x: x[1], reverse=True)
    # Filter numeric-only tokens & very short tokens
    result = [t for t, s in scored if len(t) > 2 and not t.isdigit()]
    return result[:top_k]


def keyword_alignment(resume_text: str, jd_text: str) -> Tuple[List[str], List[str], List[Tuple[str, str, int]]]:
    resume_kws = set(tfidf_keywords(resume_text, top_k=80))
    jd_kws = set(tfidf_keywords(jd_text, top_k=120))

    # Fuzzy match similar tokens to reduce duplicates (e.g., "machine learning" vs "ml")
    aligned: set = set()
    for r in list(resume_kws):
        for j in list(jd_kws):
            if fuzz.partial_ratio(r, j) >= 90:
                aligned.add(j)

    missing = sorted([kw for kw in jd_kws if kw not in aligned])
    present = sorted([kw for kw in jd_kws if kw in aligned])

    # Suggestions for synonyms (pair a missing JD kw with closest resume kw)
    suggestions: List[Tuple[str, str, int]] = []
    for m in missing:
        best_kw = None
        best = 0
        for r in resume_kws:
            score = fuzz.token_set_ratio(m, r)
            if score > best:
                best = score
                best_kw = r
        if best_kw and best >= 60:
            suggestions.append((m, best_kw, best))

    return present, missing, suggestions


def craft_summary(jd_text: str, present_kws: List[str], name_hint: str = "Candidate") -> str:
    # A concise, ATS-safe professional summary tailored to the JD and keywords found
    focus = ", ".join(present_kws[:6]) if present_kws else "role-aligned skills"
    return (
        f"{name_hint} — results-driven professional with proven experience aligning to the target role. "
        f"Strengths include {focus}. Adept at translating business goals into deliverables, collaborating cross-functionally, "
        f"and continuously improving processes to drive measurable outcomes."
    )


def normalize_header(h: str) -> str:
    h = h.strip().lower()
    mapping = {
        'professional summary': 'Summary', 'summary': 'Summary', 'profile': 'Summary', 'objective': 'Summary',
        'skills': 'Skills', 'key skills': 'Skills', 'technical skills': 'Skills', 'core competencies': 'Skills',
        'experience': 'Experience', 'work experience': 'Experience', 'professional experience': 'Experience', 'employment': 'Experience',
        'projects': 'Projects', 'relevant projects': 'Projects', 'academic projects': 'Projects',
        'education': 'Education',
        'certifications': 'Certifications',
        'achievements': 'Achievements',
        'publications': 'Publications'
    }
    return mapping.get(h, h.title())


def inject_keywords_into_skills(skills_lines: List[str], missing_kws: List[str]) -> List[str]:
    # Keep skills concise—add up to 10 missing keywords that look like hard skills
    hard_skill_like = [kw for kw in missing_kws if re.search(r"[a-z]", kw) and len(kw) <= 25]
    to_add = []
    for kw in hard_skill_like:
        # Avoid near-duplicates
        if all(fuzz.token_set_ratio(kw, s) < 90 for s in skills_lines):
            to_add.append(kw)
        if len(to_add) >= 10:
            break
    if to_add:
        return skills_lines + [", ".join(to_add)]
    return skills_lines


def reorder_experience_by_relevance(exp_lines: List[str], jd_kws: List[str]) -> List[str]:
    # Score each bullet/line by overlap with JD keywords
    def score(line: str) -> int:
        return max((fuzz.partial_ratio(line.lower(), kw) for kw in jd_kws), default=0)
    return sorted(exp_lines, key=score, reverse=True)


def build_docx(parsed: ParsedResume, jd_text: str, present: List[str], missing: List[str]) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx is not installed.")

    jd_kws = tfidf_keywords(jd_text, top_k=100)

    doc = Document()

    # Set a clean ATS-friendly base style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Contact block
    if parsed.contact_block:
        p = doc.add_paragraph(parsed.contact_block)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title
    h = doc.add_heading('Resume', level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary
    summary = craft_summary(jd_text, present)
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)

    # Skills
    skills_lines = []
    for k, v in parsed.sections.items():
        if normalize_header(k) == 'Skills':
            skills_lines.extend(v)
    skills_lines = [re.sub(r"^[\-•\*]\s*", "", s) for s in skills_lines]
    skills_lines = [s for s in skills_lines if len(s.split()) <= 40]
    skills_lines = inject_keywords_into_skills(skills_lines, missing)

    if skills_lines:
        doc.add_heading('Skills', level=1)
        for s in skills_lines:
            if ',' in s and len(s) < 120:  # compact line of comma-separated skills
                doc.add_paragraph(s)
            else:
                doc.add_paragraph(s, style=None)

    # Experience (reordered by JD relevance)
    exp_lines = []
    for k, v in parsed.sections.items():
        if normalize_header(k) == 'Experience':
            exp_lines.extend(v)
    if exp_lines:
        doc.add_heading('Experience', level=1)
        for line in reorder_experience_by_relevance(exp_lines, jd_kws):
            if BULLET_PATTERN.match(line):
                doc.add_paragraph(BULLET_PATTERN.sub("", line), style=None)
            else:
                doc.add_paragraph(line, style=None)

    # Projects
    projects = []
    for k, v in parsed.sections.items():
        if normalize_header(k) == 'Projects':
            projects.extend(v)
    if projects:
        doc.add_heading('Projects', level=1)
        for ptxt in projects:
            doc.add_paragraph(BULLET_PATTERN.sub("", ptxt))

    # Education
    edu = []
    for k, v in parsed.sections.items():
        if normalize_header(k) == 'Education':
            edu.extend(v)
    if edu:
        doc.add_heading('Education', level=1)
        for e in edu:
            doc.add_paragraph(BULLET_PATTERN.sub("", e))

    # Certifications (optional)
    certs = []
    for k, v in parsed.sections.items():
        if normalize_header(k) == 'Certifications':
            certs.extend(v)
    if certs:
        doc.add_heading('Certifications', level=1)
        for c in certs:
            doc.add_paragraph(BULLET_PATTERN.sub("", c))

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def ats_checks(text: str) -> Dict[str, str]:
    issues = {}
    # Check length
    words = len(text.split())
    if words < 250:
        issues['Length'] = f"Resume seems short ({words} words). Consider adding impact bullets."
    elif words > 1200:
        issues['Length'] = f"Resume seems long ({words} words). Consider trimming to 1–2 pages."

    # Contact info presence
    if not CONTACT_PATTERN.search(text):
        issues['Contact'] = "Could not detect email/phone/LinkedIn/GitHub. Add a contact line at top."

    # Tables/images (heuristic):
    if ' | ' in text:
        issues['Tables'] = "Detected vertical bars that might be table remnants. Ensure plain text layout."

    # Fancy symbols
    if re.search(r"[■◆●◦★]", text):
        issues['Symbols'] = "Replace special symbols with simple '-' bullets for ATS."

    return issues


def build_change_log(present: List[str], missing: List[str], suggestions: List[Tuple[str, str, int]], issues: Dict[str, str]) -> str:
    lines = []
    lines.append("Optimisation Summary (ATS & Keyword Alignment)\n" + "="*50)
    lines.append(f"Matched JD Keywords: {len(present)} | Missing (addressed in Skills/Summary): {len(missing)}")
    if present:
        lines.append("\nTop Matched Keywords:\n- " + "\n- ".join(present[:12]))
    if missing:
        lines.append("\nHigh-Value Missing Keywords Considered:\n- " + "\n- ".join(missing[:12]))
    if suggestions:
        sug_fmt = [f"'{m}' ~ '{r}' (sim: {s})" for m, r, s in suggestions[:10]]
        lines.append("\nSynonym/Alignment Suggestions:\n- " + "\n- ".join(sug_fmt))

    if issues:
        lines.append("\nATS Formatting Checks:")
        for k, v in issues.items():
            lines.append(f"- {k}: {v}")

    lines.append("\nStructural Changes Applied:\n- Added/updated 'Summary' tailored to JD\n- Normalised headers (Summary, Skills, Experience, Projects, Education)\n- Reordered Experience bullets by JD relevance\n- Injected up to 10 missing hard-skill keywords into Skills\n- Flattened any table-like text for ATS compatibility")

    return "\n".join(lines)


# ----------------------------- Streamlit UI ----------------------------- #
st.set_page_config(page_title="AI Resume Optimiser", page_icon="🧠", layout="wide")
st.title("🧠 AI Resume Optimiser & Generator (ATS-friendly)")
st.caption("Upload your resume and paste the target Job Description to generate an optimised, ATS-ready resume and a concise change log.")

with st.sidebar:
    st.header("Inputs")
    resume_file = st.file_uploader("Upload Resume (PDF or DOCX)", type=["pdf", "docx"], accept_multiple_files=False)
    jd_source = st.radio("Job Description Source", ["Paste Text", "Upload Text File"], horizontal=True)
    jd_text = ""
    if jd_source == "Paste Text":
        jd_text = st.text_area("Paste the Job Description (JD)", height=240, placeholder="Paste the full JD here…")
    else:
        jd_file = st.file_uploader("Upload JD (.txt)", type=["txt"], accept_multiple_files=False)
        if jd_file is not None:
            jd_text = jd_file.read().decode("utf-8", errors="ignore")

    run_btn = st.button("🔧 Optimise Resume", use_container_width=True)

# Main panels
col1, col2 = st.columns([1,1])

if run_btn:
    if not resume_file or not jd_text.strip():
        st.error("Please upload a resume and provide the job description text.")
        st.stop()

    # Read resume
    with st.spinner("Parsing resume…"):
        bytes_data = resume_file.read()
        if resume_file.type == 'application/pdf' or resume_file.name.lower().endswith('.pdf'):
            resume_text = read_pdf(bytes_data)
        else:
            resume_text = read_docx(bytes_data)

        lines = clean_lines(resume_text)
        parsed = split_sections(lines)

    with st.spinner("Analysing keywords & ATS checks…"):
        present, missing, suggestions = keyword_alignment(parsed.raw_text, jd_text)
        issues = ats_checks(parsed.raw_text)
        change_log = build_change_log(present, missing, suggestions, issues)

    with st.spinner("Generating optimised DOCX…"):
        try:
            docx_bytes = build_docx(parsed, jd_text, present, missing)
        except Exception as e:
            st.exception(e)
            st.stop()

    with col1:
        st.subheader("Optimised Resume (Download)")
        out_name = re.sub(r"\W+", "_", resume_file.name.rsplit('.', 1)[0]) + "_optimised.docx"
        st.download_button(
            label="⬇️ Download Optimised Resume (DOCX)",
            data=docx_bytes,
            file_name=out_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

        st.divider()
        st.subheader("Keyword Match Overview")
        st.metric("Matched JD Keywords", len(present))
        st.metric("Missing (considered)", len(missing))
        if present:
            st.write("**Top matched:**")
            st.write(", ".join(present[:20]))
        if missing:
            st.write("**High-value missing considered:**")
            st.write(", ".join(missing[:20]))

    with col2:
        st.subheader("What changed & why (ATS log)")
        st.text(change_log)

    st.success("Done! You can download the optimised resume and review the change log above.")

else:
    with col1:
        st.info("👈 Upload your resume and provide the JD to begin.")
    with col2:
        st.write("**Tips for higher ATS scores:**")
        st.markdown(
            "- Keep layout simple (no tables/images).\n"
            "- Use clear section headers: Summary, Skills, Experience, Projects, Education.\n"
            "- Mirror the employer's phrasing for hard skills and tools (but stay truthful).\n"
            "- Start bullets with strong verbs and include metrics (e.g., 'reduced cost by 18%').\n"
        )
